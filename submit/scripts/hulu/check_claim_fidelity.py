#!/usr/bin/env python3
"""Claim fidelity — does the cited work actually say what the manuscript says it says?

WHAT THIS ADDS TO `verify_refs.py`

`verify_refs.py` answers "does this reference exist, and are its authors who you say they
are". It cannot answer the question a reader actually asks, which is whether the *sentence
citing it* is true of it. A citation can be perfectly real and still be attached to a claim
the source never makes. That failure survives every existing gate: the DOI resolves, the
authors match, the reference list renders, and the sentence is wrong.

The incident this was built from: a manuscript sentence read "the field has begun to offer
the chair [41]". The cited work uses the word "chair" zero times; it uses "advocate" four
times, twice in the sense the manuscript was reaching for. A co-author caught it by reading
the source. Nothing in the toolkit could have.

THREE PROBES, ORDERED BY HOW CHECKABLE THE CLAIM IS

  1. QUOTED TEXT (`CITED_QUOTE_ABSENT`, major)
     A quotation attributed to a source is the one claim with an exact answer: the words are
     in that document or they are not. This is where a quote-operator inversion lives — a
     manuscript that quotes a source as saying "only if" where the source said "so long as"
     has reversed its meaning while looking verbatim.

  2. CONCEPT ATTRIBUTION (`ATTRIBUTION_UNSUPPORTED`, prompt)
     "X proposes <thing> [12]". Paraphrase is legitimate and ordinary, so this cannot be a
     blocker and cannot fire on mere wording differences. It fires only in the extreme case:
     NOT ONE content word of the attributed span appears anywhere in the source, in any
     morphological form. A real paraphrase almost always keeps at least one content stem
     ("propose a framework for oversight" keeps "oversight"); an attribution to the wrong
     concept keeps none. That is the whole precision argument, and it is why the verdict is
     phrased as a prompt to go read the source rather than as a defect.

  3. CARDINAL CLAIMS (`ORDINAL_CLAIM_UNSUPPORTED`, prompt)
     "reports three strategies [12]" is cheap to get wrong and cheap to check. Fires only
     when the noun IS present in the source (so we know the concept was located) but the
     stated cardinal never appears near it. If the noun is absent entirely that is probe 2's
     business, not this one.

DELIBERATELY NOT CHECKED: float ordinals ("as in Table 2 of [12]"). A manuscript refers to
its OWN Table 2 constantly, and often in a sentence that also carries a citation, so the
probe would fire mostly on correct prose. A check that is right in principle and wrong in
practice is a check that gets switched off, and takes the honest ones with it.

KNOWN LIMITATION, STATED SO IT IS NOT MISTAKEN FOR COVERAGE: a *small* alteration inside a
long quotation — flipping "so long as" to "only if", two tokens in twenty — lands inside the
20% loss that `_quote_match` tolerates as extraction damage, so it is reported as
`CITED_QUOTE_UNRESOLVED` (a prompt) rather than as a fabrication. That is not a bug to fix by
tightening the threshold: nothing in a single extracted text distinguishes "the author changed
two words" from "the extractor dropped two words". The claim is still surfaced for a human;
only the severity is withheld. A heavier rewrite falls outside the tolerance and is major. Both
sides of that boundary are pinned in the challenge card.

WHY ABSENCE IS ONLY EVIDENCE FROM A WHOLE DOCUMENT

Every verdict here is an argument from absence, so it is only as good as the text it
searched. Two guards make that argument honest:

  * The source text must be substantial (`MIN_SOURCE_TOKENS`). A PDF that converted to its
    abstract alone would make every probe scream; such a source is reported as unusable
    (`sources_too_short`) and produces no findings. A quote graded absent against a short
    source is downgraded to unresolved, never to a defect.
  * Quote matching goes through `_quote_match.py`, so a correct quote read through a dirty
    extraction (a bled reference column, PDF line numbers, a hyphen split across a line) is
    reported as UNRESOLVED rather than as a fabrication. That module exists because a
    contiguous-substring test produced thirteen false absences in one day.

INPUT CONTRACT

Full texts are supplied as a directory of ALREADY-CONVERTED text — this detector never
fetches. `/fulltext-retrieval` produces exactly that (`fetch_oa.py` writes
`<doi-with-unsafe-chars-underscored>.pdf`, `pdf_to_md.py` converts it alongside), and a
detector that reached the network would be neither deterministic nor CI-runnable.

A citation is resolved to a file by, in order: an explicit `--refmap`; the DOI from the
`--bib` entry for `[@key]`; the DOI parsed out of a numbered reference list in the
manuscript for `[N]`; or a file whose stem is the citation key itself. Citations that
resolve to nothing are counted, never guessed at, and never produce a finding.

Usage:
    check_claim_fidelity.py --manuscript m.md --fulltext-dir fulltext/ [--bib refs.bib]
    check_claim_fidelity.py --manuscript m.md --fulltext-dir ft/ --out qc/claim_fidelity.json --strict

Exit 0 when no major verdict fires (prompts alone do not fail). With --strict, exit 1 if any
major fires. Stdlib only; .docx read via python-docx when available.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _quote_match import match_quality, normalize, tokens  # noqa: E402  (vendored, same-dir)

DETECTOR = "check_claim_fidelity"

# A source shorter than this is an extraction stub (an abstract, a landing page, a paywall
# interstitial). Absence proves nothing against it, so it produces no findings at all.
MIN_SOURCE_TOKENS = 400
# A quoted span shorter than this is not a quotation, it is a scare-quoted term.
MIN_QUOTE_TOKENS = 6
# How far after a closing quote a citation may sit and still be its attribution.
CITE_WINDOW_CHARS = 120
# Attributed span: tokens between the attribution verb and the citation that are examined.
MAX_SPAN_TOKENS = 14
# A stem must be this long before a prefix match is meaningful ("dat" matches too much).
MIN_STEM = 4
# How far from an occurrence of the noun the source may state the cardinal.
CARDINAL_WINDOW = 15

# Verbs that assign a claim to a cited work. The citation in the same sentence is the real
# filter; this list only has to be broad enough to find the point where attribution starts.
ATTRIB_VERB = re.compile(
    r"\b("
    r"propos\w*|offer\w*|introduc\w*|argu\w*|report\w*|show\w*|demonstrat\w*|describ\w*|"
    r"suggest\w*|conclud\w*|observ\w*|advocat\w*|recommend\w*|claim\w*|assert\w*|"
    r"present\w*|identif\w*|characteri[sz]\w*|defin\w*|coin\w*|"
    r"found|finds|note[ds]?|term(?:s|ed)|call(?:s|ed)|frame[sd]?"
    r")\b",
    re.IGNORECASE,
)

CITE_KEY = re.compile(r"\[([^\]]*@[^\]]+)\]")
CITE_NUM = re.compile(r"\[(\d{1,3}(?:\s*[-–,;]\s*\d{1,3})*)\]")
DOI_RE = re.compile(r"\b10\.\d{4,9}/[^\s\"<>,;\]]+", re.IGNORECASE)
QUOTE_RE = re.compile(r"[\"“]([^\"“”]{25,700})[\"”]")

NUM_WORDS = {
    "one": 1, "two": 2, "three": 3, "four": 4, "five": 5, "six": 6,
    "seven": 7, "eight": 8, "nine": 9, "ten": 10, "eleven": 11, "twelve": 12,
}
CARDINAL_CLAIM = re.compile(
    r"\b(one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|\d{1,3})\s+"
    r"((?:[a-z]+[- ]){0,2}[a-z]{3,}s)\b",
    re.IGNORECASE,
)

# Function words plus the vocabulary every paper shares. A span made only of these carries no
# checkable content, so it is skipped rather than flagged.
STOPWORDS = frozenset("""
a an the this that these those there here it its it's they them their theirs we our us you your
and or but nor for so yet as if then than because while whereas although though however
of in on at by to from with within without into onto over under between among across during
is are was were be been being being has have had having do does did done
can could may might must shall should will would
not no none nor only just also both each either neither all any some many much more most few less least
such same other others another new recent previous prior earlier later further additional
which who whom whose what when where why how
study studies paper papers work works article articles author authors research researchers
result results finding findings data dataset datasets analysis analyses approach approaches
method methods methodology model models framework frameworks evidence report reports
review reviews literature field fields context setting settings case cases example examples
use used using uses need needs based given note noted show shown seen make made take taken
group groups patient patients subject subjects participant participants sample samples
value values level levels rate rates number numbers effect effects
""".split())


# --------------------------------------------------------------------------- reading input

def read_text(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        try:
            import docx  # type: ignore
        except ImportError:
            print(f"ERROR: reading {path.name} needs python-docx (pip install python-docx)",
                  file=sys.stderr)
            raise SystemExit(2)
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                parts += [c.text for c in row.cells]
        return "\n".join(parts)
    return path.read_text(encoding="utf-8", errors="replace")


def strip_noise(md: str) -> str:
    """Remove regions whose text is not manuscript prose making claims.

    Fenced code, reviewer/quoted blockquote lines and table rows are dropped; the reference
    list is truncated away because a numbered bibliography would feed probe 3 a cardinal per
    line ("12 Smith J, 4 Brown K") and probe 2 a citation per line.
    """
    md = re.sub(r"```.*?```", " ", md, flags=re.S)
    md = re.sub(r"^\s*>.*$", " ", md, flags=re.M)
    md = re.sub(r"^\s*\|.*$", " ", md, flags=re.M)
    cut = re.search(r"^#{1,3}\s*\**\s*(References|REFERENCES|Bibliography|Works Cited)\b",
                    md, flags=re.M)
    return md[:cut.start()] if cut else md


def reference_section(md: str) -> str:
    m = re.search(r"^#{1,3}\s*\**\s*(References|REFERENCES|Bibliography|Works Cited)\b",
                  md, flags=re.M)
    return md[m.start():] if m else ""


# ------------------------------------------------------------------- citation -> source file

def parse_bib(path: Path) -> dict[str, str]:
    """citation key -> DOI, for the entries that carry one."""
    out: dict[str, str] = {}
    text = path.read_text(encoding="utf-8", errors="replace")
    for chunk in re.split(r"\n(?=@)", text):
        key_m = re.match(r"@\w+\s*\{\s*([^,\s]+)\s*,", chunk)
        if not key_m:
            continue
        doi_m = re.search(r"\bdoi\s*=\s*[{\"]\s*(?:https?://doi\.org/)?([^}\"\s]+)",
                          chunk, re.IGNORECASE)
        if doi_m:
            out[key_m.group(1)] = doi_m.group(1).rstrip(".,;")
    return out


REF_ENTRY_START = re.compile(r"^\s*\[?(\d{1,3})[.)\]]\s+\S")


def parse_numbered_refs(ref_text: str) -> dict[str, str]:
    """reference number -> DOI, read off a numbered reference list in the manuscript.

    Entries are accumulated as BLOCKS rather than lines: a rendered reference wraps, and the DOI
    is usually on the continuation line, so a line-scoped read finds the number and the DOI in
    different rows and resolves nothing.
    """
    out: dict[str, str] = {}
    num: str | None = None
    buf: list[str] = []

    def flush() -> None:
        if num is None:
            return
        doi = DOI_RE.search(" ".join(buf))
        if doi:
            out[num] = doi.group(0).rstrip(".,;")

    for line in ref_text.splitlines():
        m = REF_ENTRY_START.match(line)
        if m:
            flush()
            num, buf = m.group(1), [line]
        elif num is not None:
            if not line.strip():
                flush()
                num, buf = None, []
            else:
                buf.append(line)
    flush()
    return out


def safe_doi_name(doi: str) -> str:
    """Filesystem-safe stem for a DOI — identical to /fulltext-retrieval fetch_oa.py, so the
    files that skill writes are found without any hand-built map."""
    return re.sub(r"[^\w\-.]", "_", doi)


def index_fulltext(d: Path) -> dict[str, Path]:
    """lowercased stem -> file, over the text formats pdf_to_md.py and friends emit."""
    idx: dict[str, Path] = {}
    for p in sorted(d.rglob("*")):
        if p.is_file() and p.suffix.lower() in (".md", ".txt", ".markdown", ".text"):
            idx.setdefault(p.stem.lower(), p)
    return idx


def expand_numeric(raw: str) -> list[str]:
    """"4-11, 15" -> ['4'..'11', '15']. Ranges must expand or a cited number inside one reads
    as uncited — the same trap that produced false gaps in the citation-order work."""
    out: list[str] = []
    for part in re.split(r"\s*[,;]\s*", raw):
        rng = re.match(r"^(\d{1,3})\s*[-–]\s*(\d{1,3})$", part.strip())
        if rng:
            lo, hi = int(rng.group(1)), int(rng.group(2))
            if lo <= hi and hi - lo <= 60:
                out += [str(n) for n in range(lo, hi + 1)]
                continue
        p = part.strip()
        if p.isdigit():
            out.append(p)
    return out


class Resolver:
    """Turns the citation tokens in a sentence into readable source texts."""

    def __init__(self, fulltext: dict[str, Path], key_doi: dict[str, str],
                 num_doi: dict[str, str], refmap: dict[str, str]):
        self.fulltext = fulltext
        self.key_doi = key_doi
        self.num_doi = num_doi
        self.refmap = refmap
        self.unresolved: set[str] = set()
        self._cache: dict[str, tuple[Path, str] | None] = {}

    def _file_for(self, token: str) -> Path | None:
        mapped = self.refmap.get(token)
        candidates: list[str] = []
        if mapped:
            candidates += [safe_doi_name(mapped), mapped]
        doi = self.key_doi.get(token) or self.num_doi.get(token)
        if doi:
            candidates.append(safe_doi_name(doi))
        candidates.append(token)
        for c in candidates:
            hit = self.fulltext.get(c.lower())
            if hit:
                return hit
        return None

    def resolve(self, token: str) -> tuple[Path, str] | None:
        if token in self._cache:
            return self._cache[token]
        path = self._file_for(token)
        got = (path, read_text(path)) if path else None
        if got is None:
            self.unresolved.add(token)
        self._cache[token] = got
        return got

    @staticmethod
    def citations_in(text: str) -> list[str]:
        out: list[str] = []
        for m in CITE_KEY.finditer(text):
            out += [k.strip().lstrip("@").strip()
                    for k in re.split(r"[;,]", m.group(1)) if "@" in k]
        for m in CITE_NUM.finditer(text):
            out += expand_numeric(m.group(1))
        return [t for t in out if t]


# ------------------------------------------------------------------------- text utilities

def stem(word: str) -> str:
    """Crude suffix strip — enough to let 'strategies' find 'strategy' by prefix."""
    w = word.lower()
    for suf, repl in (("ies", "i"), ("ing", ""), ("ies", "y"), ("es", ""), ("ed", ""), ("s", "")):
        if w.endswith(suf) and len(w) - len(suf) >= MIN_STEM:
            return w[: len(w) - len(suf)] + repl
    return w


class Source:
    """A resolved full text, prepared once for repeated membership questions."""

    def __init__(self, path: Path, raw: str):
        self.path = path
        self.raw = raw
        self.tokens = tokens(raw)
        self.token_set = frozenset(self.tokens)
        self.prefixes = frozenset(t[:MIN_STEM] for t in self.tokens if len(t) >= MIN_STEM)
        self.usable = len(self.tokens) >= MIN_SOURCE_TOKENS

    def has_term(self, word: str) -> bool:
        """Morphology-tolerant presence: exact token, or any source token sharing this word's
        stem prefix and extending it plausibly."""
        w = word.lower()
        if w in self.token_set:
            return True
        s = stem(w)
        if len(s) < MIN_STEM:
            return False
        if s in self.token_set:
            return True
        if s[:MIN_STEM] not in self.prefixes:
            return False
        return any(t.startswith(s) and len(t) - len(s) <= 3 for t in self.token_set)

    def positions(self, word: str) -> list[int]:
        s = stem(word.lower())
        if len(s) < MIN_STEM:
            return []
        return [i for i, t in enumerate(self.tokens) if t.startswith(s) and len(t) - len(s) <= 3]


ABBREV = re.compile(r"\b(?:et al|e\.g|i\.e|cf|vs|Fig|Figs|Tab|No|Dr|Prof|approx)\.$",
                    re.IGNORECASE)


def sentences(text: str) -> list[str]:
    out: list[str] = []
    buf = ""
    for piece in re.split(r"(?<=[.!?])\s+", text):
        buf = f"{buf} {piece}".strip() if buf else piece
        if ABBREV.search(buf.strip()):
            continue
        out.append(buf)
        buf = ""
    if buf:
        out.append(buf)
    return [s.strip() for s in out if s.strip()]


def content_words(span: str) -> list[str]:
    ws = re.findall(r"[A-Za-z][A-Za-z'-]{2,}", span)
    return [w for w in ws if w.lower() not in STOPWORDS and len(w) >= MIN_STEM]


# ------------------------------------------------------------------------------ the probes

def probe_quotes(md: str, res: Resolver, findings: list[dict]) -> tuple[int, list[tuple[int, int]]]:
    """Quoted text attributed to a source must be in that source. Returns (checked, spans)."""
    checked = 0
    spans: list[tuple[int, int]] = []
    for m in QUOTE_RE.finditer(md):
        quote = m.group(1).strip()
        if len(tokens(quote)) < MIN_QUOTE_TOKENS:
            continue
        tail = md[m.end(): m.end() + CITE_WINDOW_CHARS]
        cites = Resolver.citations_in(tail)
        if not cites:
            continue
        spans.append((m.start(), m.end()))
        for token in cites[:3]:
            got = res.resolve(token)
            if not got:
                continue
            src = Source(*got)
            checked += 1
            grade = match_quality(quote, src.raw)
            if grade["grade"] in ("EXACT", "INTERLEAVED"):
                continue
            absent = grade["grade"] == "ABSENT" and src.usable
            findings.append({
                "verdict": "CITED_QUOTE_ABSENT" if absent else "CITED_QUOTE_UNRESOLVED",
                "severity": "major" if absent else "minor",
                "citation": token,
                "source": src.path.name,
                "quote": quote[:200],
                "match_grade": grade["grade"],
                "coverage": grade["coverage"],
                "message": (
                    f"Quoted text attributed to [{token}] is not in {src.path.name} in any "
                    f"reading order (coverage {grade['coverage']}). Read the source before "
                    f"keeping the quotation marks."
                ) if absent else (
                    f"Quoted text attributed to [{token}] matched {src.path.name} only as "
                    f"{grade['grade']} (coverage {grade['coverage']})"
                    + ("" if src.usable else "; the extracted source is too short to judge")
                    + " — verify by eye rather than treating this as a defect."
                ),
            })
            break
    return checked, spans


def probe_attribution(md: str, res: Resolver, quote_spans: list[tuple[int, int]],
                      findings: list[dict]) -> int:
    """A claim assigned to a source whose every content word is absent from that source."""
    checked = 0
    offset = 0
    for sent in sentences(md):
        idx = md.find(sent, offset)
        if idx >= 0:
            offset = idx + len(sent)
            # A sentence that carries a quotation is probe 1's — attributing the whole quoted
            # span to probe 2 as well would report one claim twice, under two verdicts.
            if any(idx < b and a < idx + len(sent) for a, b in quote_spans):
                continue
        cite_m = CITE_KEY.search(sent) or CITE_NUM.search(sent)
        if not cite_m:
            continue
        verbs = [v for v in ATTRIB_VERB.finditer(sent) if v.end() <= cite_m.start()]
        if not verbs:
            continue
        span = sent[verbs[-1].end(): cite_m.start()]
        words = content_words(span)[:MAX_SPAN_TOKENS]
        if not words:
            continue
        for token in Resolver.citations_in(cite_m.group(0))[:2]:
            got = res.resolve(token)
            if not got:
                continue
            src = Source(*got)
            if not src.usable:
                continue
            checked += 1
            if any(src.has_term(w) for w in words):
                continue
            findings.append({
                "verdict": "ATTRIBUTION_UNSUPPORTED",
                "severity": "minor",
                "citation": token,
                "source": src.path.name,
                "attributed_verb": verbs[-1].group(0),
                "attributed_terms": words,
                "sentence": sent.strip()[:300],
                "message": (
                    f"Nothing attributed to [{token}] here — {', '.join(words)} — appears "
                    f"anywhere in {src.path.name}, in any form. Paraphrase normally keeps at "
                    f"least one of the source's own terms; none survived. Read the source and "
                    f"confirm this is the claim it makes."
                ),
            })
            break
    return checked


def probe_cardinal(md: str, res: Resolver, findings: list[dict]) -> int:
    """"reports three strategies [12]" where the source names that noun but never that count."""
    checked = 0
    for sent in sentences(md):
        cite_m = CITE_KEY.search(sent) or CITE_NUM.search(sent)
        if not cite_m:
            continue
        if not any(v.end() <= cite_m.start() for v in ATTRIB_VERB.finditer(sent)):
            continue
        claims = [c for c in CARDINAL_CLAIM.finditer(sent[: cite_m.start()])
                  if c.group(2).split()[-1].lower() not in STOPWORDS]
        if not claims:
            continue
        for token in Resolver.citations_in(cite_m.group(0))[:2]:
            got = res.resolve(token)
            if not got:
                continue
            src = Source(*got)
            if not src.usable:
                continue
            for c in claims:
                raw_num, noun_phrase = c.group(1).lower(), c.group(2)
                noun = noun_phrase.split()[-1].split("-")[-1]
                value = NUM_WORDS.get(raw_num, None)
                if value is None:
                    if not raw_num.isdigit():
                        continue
                    value = int(raw_num)
                forms = {str(value)} | {w for w, v in NUM_WORDS.items() if v == value}
                spots = src.positions(noun)
                if not spots:
                    continue  # concept not located: probe 2's territory, not this one
                checked += 1
                near = any(
                    src.tokens[j] in forms
                    for i in spots
                    for j in range(max(0, i - CARDINAL_WINDOW),
                                   min(len(src.tokens), i + CARDINAL_WINDOW + 1))
                )
                if near:
                    continue
                findings.append({
                    "verdict": "ORDINAL_CLAIM_UNSUPPORTED",
                    "severity": "minor",
                    "citation": token,
                    "source": src.path.name,
                    "claimed": f"{raw_num} {noun_phrase}",
                    "sentence": sent.strip()[:300],
                    "message": (
                        f"This attributes \"{raw_num} {noun_phrase}\" to [{token}], but "
                        f"{src.path.name} discusses \"{noun}\" without ever stating that count "
                        f"nearby. Count them in the source."
                    ),
                })
            break
    return checked


# ------------------------------------------------------------------------------------ main

def build_report(manuscript: Path, fulltext_dir: Path, bib: Path | None,
                 refmap: dict[str, str]) -> dict:
    raw = read_text(manuscript)
    body = strip_noise(raw)
    res = Resolver(
        fulltext=index_fulltext(fulltext_dir),
        key_doi=parse_bib(bib) if bib else {},
        num_doi=parse_numbered_refs(reference_section(raw)),
        refmap=refmap,
    )

    findings: list[dict] = []
    n_quote, spans = probe_quotes(body, res, findings)
    n_attr = probe_attribution(body, res, spans, findings)
    n_card = probe_cardinal(body, res, findings)

    resolved = {t: v for t, v in res._cache.items() if v}
    short = sorted({v[0].name for t, v in resolved.items()
                    if len(tokens(v[1])) < MIN_SOURCE_TOKENS})
    return {
        "detector": DETECTOR,
        "manuscript": str(manuscript),
        "fulltext_dir": str(fulltext_dir),
        "sources_resolved": len({v[0] for v in resolved.values()}),
        "sources_unresolved": sorted(res.unresolved),
        "sources_too_short": short,
        "claims_checked": {"quotes": n_quote, "attributions": n_attr, "cardinals": n_card},
        "findings": findings,
    }


def main() -> int:
    ap = argparse.ArgumentParser(description="Verify what a manuscript says a cited work says.")
    ap.add_argument("--manuscript", required=True, type=Path)
    ap.add_argument("--fulltext-dir", required=True, type=Path,
                    help="directory of ALREADY-CONVERTED source texts (see /fulltext-retrieval)")
    ap.add_argument("--bib", type=Path, help="refs.bib, for [@key] -> DOI resolution")
    ap.add_argument("--refmap", type=Path,
                    help='JSON {"41": "10.1000/xyz"} for citations neither bib nor list resolves')
    ap.add_argument("--out", type=Path, help="write the JSON report here")
    ap.add_argument("--strict", action="store_true", help="exit 1 if any major verdict fires")
    args = ap.parse_args()

    if not args.manuscript.is_file():
        print(f"ERROR: manuscript not found: {args.manuscript}", file=sys.stderr)
        return 2
    if not args.fulltext_dir.is_dir():
        print(f"ERROR: --fulltext-dir not a directory: {args.fulltext_dir}", file=sys.stderr)
        return 2
    refmap: dict[str, str] = {}
    if args.refmap:
        refmap = {str(k): str(v)
                  for k, v in json.loads(args.refmap.read_text(encoding="utf-8")).items()}

    report = build_report(args.manuscript, args.fulltext_dir, args.bib, refmap)

    if args.out:
        args.out.parent.mkdir(parents=True, exist_ok=True)
        args.out.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")

    c = report["claims_checked"]
    print(f"{DETECTOR}: {report['sources_resolved']} source(s) resolved; "
          f"checked {c['quotes']} quote(s), {c['attributions']} attribution(s), "
          f"{c['cardinals']} cardinal claim(s)")
    if report["sources_unresolved"]:
        print(f"  not resolved to a full text (not checked): "
              f"{', '.join(report['sources_unresolved'][:12])}")
    if report["sources_too_short"]:
        print(f"  extracted text too short to judge absence: "
              f"{', '.join(report['sources_too_short'][:8])}")

    majors = [f for f in report["findings"] if f["severity"] == "major"]
    if not report["findings"]:
        print("OK: every checkable claim is supported by its cited source.")
        return 0
    for f in report["findings"]:
        print(f"  [{f['severity']}] {f['verdict']}: {f['message']}")
    return 1 if (args.strict and majors) else 0


if __name__ == "__main__":
    sys.exit(main())
